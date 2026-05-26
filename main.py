import streamlit as st
from docx import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from sentence_transformers import CrossEncoder
import tempfile
import os

st.title("我的AI知识助手")

uploaded_file = st.file_uploader(
    "上传文件",
    type=["pdf", "docx"]
)

if uploaded_file:

    suffix = uploaded_file.name.split(".")[-1]

    with tempfile.NamedTemporaryFile(
        delete=False,
        suffix="." + suffix
    ) as tmp:

        tmp.write(uploaded_file.read())
        temp_path = tmp.name

    text = ""

    try:

        # PDF读取
        if suffix == "pdf":

            loader = PyPDFLoader(temp_path)

            docs = loader.load()

            text = "\n".join(
                [doc.page_content for doc in docs]
            )

        # Word读取
        elif suffix == "docx":

            doc = Document(temp_path)

            text = "\n".join(
                [p.text for p in doc.paragraphs]
            )

        # 显示原始内容
        st.success("文件读取成功")

        st.text_area(
            "文件内容预览",
            text,
            height=300
        )

        # 文本切块
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=250,
            chunk_overlap=50
        )

        chunks = splitter.split_text(text)

        st.success(
            f"切块成功，共 {len(chunks)} 块"
        )

        # 转文档对象
        documents = [
            Document(page_content=chunk)
            for chunk in chunks
        ]

        with st.spinner("正在加载Embedding模型..."):

            embeddings = HuggingFaceEmbeddings(
                model_name="BAAI/bge-base-zh-v1.5",
                model_kwargs={
                    "local_files_only":True   #注意第一次缓存到本地之后，后面的rerank也是，限制只在本地寻找
                }
            )

        with st.spinner("正在创建向量数据库..."):

            vector_db = FAISS.from_documents(
                documents,
                embeddings
            )

        st.success("向量数据库构建成功")

        st.write(
            f"当前文本块数量：{len(documents)}"
        )

        # 展示前5个文本块
        for i, chunk in enumerate(chunks[:5]):

            st.text_area(
                f"文本块 {i+1}",
                chunk,
                height=120
            )

        # 用户问题输入
        query = st.text_input(
            "请输入问题"
        )

        # 在提问外面加载排序模型避免重复加载
        @st.cache_resource
        def load_reranker():
            return CrossEncoder(
                "BAAI/bge-reranker-base",
                local_files_only=True
            )

        reranker = load_reranker()

        if query:

            with st.spinner("正在检索知识库..."):

                # 创建检索器  从top5中选3
                retriever = vector_db.as_retriever(
                    search_kwargs={"k":5}
                )

                # 搜索最相关文本
                docs = retriever.invoke(query)

            with st.spinner("正在进行二次排序..."):

                pairs = [
                    [query, doc.page_content]
                    for doc in docs
                ]
                #计算得分
                scores = reranker.predict(
                    pairs
                )
                # 排序
                reranked_results = sorted(
                    zip(scores, docs),
                    reverse=True,
                    key=lambda x:x[0]
                )
                #取出来前三
                top_docs = [
                    doc
                    for score, doc in reranked_results[:3]
                ]

            st.success(
                f"找到 {len(top_docs)} 条相关内容"
            )
            #展示结果
            for i, doc in enumerate(top_docs):

                st.text_area(
                    f"检索结果{i+1}",
                    doc.page_content,
                    height=150
                )

    except Exception as e:

        st.error(
            f"发生错误: {str(e)}"
        )

    finally:

        # 删除临时文件
        if os.path.exists(temp_path):
            os.remove(temp_path)